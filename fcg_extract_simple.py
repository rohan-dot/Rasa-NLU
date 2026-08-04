#!/usr/bin/env python3
"""fcg_extract_simple.py — checklist items -> FCG extraction CSV (v2).

v2 adds, per the agreed design:
  * multi-source dossiers: FCG country JSON + MEIS airfield slices +
    SVC_RMK rows + udl NOTAM blocks (via country_sources.py), each with
    its OWN char budget so FCG text can never be squeezed out
  * batched extraction: ~104 fields split into batches of FIELDS_PER_CALL
    (default 20) -> ~5-6 LLM calls per country, immune to answer-JSON
    truncation
  * resume per (country, batch): safe to Ctrl-C anytime; changing the
    items file invalidates only affected batches
  * dossier audit files in data/work/dossiers/ so every answer can be
    traced to the exact text the model saw

Usage:
  python fcg_extract_simple.py items.json --countries data/fcg_countries \
      --sources data/sources
  (omit --sources to extract from FCG JSONs alone)

Env (same vars as everything else): LITELLM_BASE_URL / LITELLM_API_KEY /
AGENT_MODEL / AGENT_TLS_VERIFY, plus optional budget knobs:
  FCG_BUDGET_FCG=9000 FCG_BUDGET_MEIS=6000 FCG_BUDGET_SVC=4000
  FCG_BUDGET_NOTAM=2000 FCG_FIELDS_PER_CALL=20 FCG_CONCURRENCY=4
"""

import argparse
import asyncio
import csv
import hashlib
import json
import os
import re
import sys
import unicodedata
from pathlib import Path

# ----------------------------------------------------------------------------
# Config (env)
# ----------------------------------------------------------------------------

def _env(*names, default=""):
    for n in names:
        v = os.environ.get(n)
        if v:
            return v
    return default


BASE_URL = _env("FCG_LLM_BASE_URL", "LITELLM_BASE_URL",
                default="http://localhost:4000/v1")
API_KEY = _env("FCG_LLM_API_KEY", "LITELLM_API_KEY", default="sk-local")
MODEL = _env("FCG_LLM_MODEL", "AGENT_MODEL", default="claude-opus-4-8")
TLS_VERIFY = _env("FCG_TLS_VERIFY", "AGENT_TLS_VERIFY",
                  default="true").strip().lower() not in ("false", "0", "no")
CONCURRENCY = int(_env("FCG_CONCURRENCY", default="4"))
MAX_TOKENS = int(_env("FCG_MAX_TOKENS", default="4000"))
RETRIES = 3

FIELDS_PER_CALL = int(_env("FCG_FIELDS_PER_CALL", default="20"))
BUDGET_FCG = int(_env("FCG_BUDGET_FCG", default="9000"))
BUDGET_MEIS = int(_env("FCG_BUDGET_MEIS", default="6000"))
BUDGET_SVC = int(_env("FCG_BUDGET_SVC", default="4000"))
BUDGET_NOTAM = int(_env("FCG_BUDGET_NOTAM", default="2000"))

CORE_FIELDS = [
    ("overflight_raw",
     "overflight permissions, diplomatic overflight clearance requirements, "
     "and overflight lead times"),
    ("diplomatic_lead_time_raw",
     "diplomatic clearance lead times for landing rights"),
    ("entry_exit_airports_raw",
     "designated entry/exit airports, airports of entry (AOE), and which "
     "airports foreign aircraft may use"),
    ("entry_exit_airports_summary",
     "a comma-separated list of every ICAO/IATA airport code mentioned as a "
     "designated entry/exit airport (codes only, no prose)"),
    ("customs_immigration_raw",
     "customs, immigration, and CIQ requirements"),
    ("airfield_restrictions_raw",
     "airfield restrictions, PPR requirements, and prohibited airfields"),
    ("operating_hours_raw", "airfield or ATC operating hours"),
    ("hazmat_raw", "HAZMAT, dangerous goods, and munitions rules"),
    ("aircard_cash_raw",
     "AIR Card acceptance, fuel payment, and cash payment rules"),
    ("country_specific_raw",
     "country-specific notes, special requirements, or restrictions not "
     "covered by other categories"),
]

QUESTION = ("Quote verbatim all dossier text about {topic}. Preserve the "
            "original wording. If the dossier says nothing about this, "
            "answer exactly NA.")

# ----------------------------------------------------------------------------
# Items -> fields
# ----------------------------------------------------------------------------

def slug(title: str) -> str:
    s = unicodedata.normalize("NFKD", title).encode("ascii", "ignore").decode()
    s = re.sub(r"[^a-z0-9]+", "_", s.lower()).strip("_")
    s = re.sub(r"_+", "_", s)[:34] or "item"
    return s if s.endswith(("_raw", "_summary")) else s + "_raw"


def load_items(path: Path) -> list[dict]:
    if path.suffix.lower() == ".docx":
        from docx import Document
        doc = Document(str(path))
        titles = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
        for table in doc.tables:
            for row in table.rows:
                cells = [c.text.strip() for c in row.cells if c.text.strip()]
                if cells:
                    titles.append(" ".join(cells))
        return [{"title": t, "notes": ""} for t in titles]
    data = json.loads(path.read_text(encoding="utf-8"))
    if isinstance(data, dict):
        data = data.get("items", [])
    out = []
    for e in data:
        if isinstance(e, str) and e.strip():
            out.append({"title": e.strip(), "notes": ""})
        elif isinstance(e, dict) and e.get("title"):
            out.append({"title": str(e["title"]).strip(),
                        "notes": str(e.get("notes", "")).strip()})
    return out


def build_fields(items: list[dict]) -> list[tuple[str, str]]:
    fields, seen = [], set()
    for name, topic in CORE_FIELDS:
        fields.append((name, QUESTION.format(topic=topic)))
        seen.add(name)
    for item in items:
        name = slug(item["title"])
        if name in seen:
            continue
        seen.add(name)
        topic = item["title"] + (f" (focus: {item['notes']})"
                                 if item["notes"] else "")
        fields.append((name, QUESTION.format(topic=topic)))
    return fields


def batch_fields(fields, size=None):
    size = size or FIELDS_PER_CALL
    return [fields[i:i + size] for i in range(0, len(fields), size)]

# ----------------------------------------------------------------------------
# Lenient JSON parsing (three-tier, truncation-safe)
# ----------------------------------------------------------------------------

_FENCE = re.compile(r"```(?:json)?\s*(.*?)\s*```", re.DOTALL)


def _balance(t: str) -> str:
    stack, in_str, esc = [], False, False
    for ch in t:
        if esc:
            esc = False
            continue
        if ch == "\\":
            esc = in_str
            continue
        if ch == '"':
            in_str = not in_str
            continue
        if in_str:
            continue
        if ch in "{[":
            stack.append(ch)
        elif ch == "}" and stack and stack[-1] == "{":
            stack.pop()
        elif ch == "]" and stack and stack[-1] == "[":
            stack.pop()
    if in_str:
        t += '"'
    for op in reversed(stack):
        t += "}" if op == "{" else "]"
    return t


def parse_lenient(text):
    if not text:
        return None
    m = _FENCE.search(text)
    cand = m.group(1) if m else text
    for attempt in (cand,
                    _balance(re.sub(r",\s*([}\]])", r"\1", cand.strip())),
                    None):
        if attempt is None:
            cut = max(cand.rfind("}"), cand.rfind(","))
            if cut <= 0:
                return None
            attempt = _balance(cand[:cut + 1].rstrip().rstrip(","))
        try:
            return json.loads(attempt)
        except json.JSONDecodeError:
            continue
    return None

# ----------------------------------------------------------------------------
# Dossier assembly (per-source budgets)
# ----------------------------------------------------------------------------

def flatten(obj, prefix="") -> list[str]:
    lines = []
    if isinstance(obj, dict):
        for k, v in obj.items():
            lines.extend(flatten(v, f"{prefix}{k}."))
    elif isinstance(obj, list):
        for i, v in enumerate(obj):
            lines.extend(flatten(v, f"{prefix}{i}."))
    else:
        t = str(obj).strip()
        if t:
            lines.append(f"{prefix[:-1]}: {t}")
    return lines


_SUFFIX = re.compile(r"_FCG[\d.]*$", re.IGNORECASE)


def country_key(path: Path) -> str:
    return _SUFFIX.sub("", path.stem).upper()


def build_dossier(country: str, fcg_path: Path, source_slice) -> str:
    try:
        data = json.loads(fcg_path.read_text(encoding="utf-8",
                                             errors="replace"))
        fcg_text = "\n".join(flatten(data))
    except json.JSONDecodeError:
        fcg_text = fcg_path.read_text(encoding="utf-8", errors="replace")

    parts = [f"=== SECTION 1: FCG (Foreign Clearance Guide) for {country} ===",
             fcg_text[:BUDGET_FCG]]

    if source_slice:
        meis = source_slice.get("meis") or []
        if meis:
            per_ap = max(400, BUDGET_MEIS // max(1, len(meis)))
            body = "\n\n".join(m["text"][:per_ap] for m in meis)
            parts += ["\n=== SECTION 2: AIRFIELD DETAILS (GIANT/MEIS report) ===",
                      body[:BUDGET_MEIS]]
        svc = source_slice.get("svc") or {}
        if svc.get("rows"):
            dr = svc.get("date_range")
            label = f" [cycle dates {dr[0]}-{dr[1]}]" if dr else ""
            parts += [f"\n=== SECTION 3: ENROUTE SUPPORT REMARKS "
                      f"(SVC_RMK){label} ===",
                      (svc.get("header", "") + "\n"
                       + "\n".join(svc["rows"]))[:BUDGET_SVC]]
        notam = source_slice.get("notam") or {}
        if notam.get("blocks"):
            dr = notam.get("date_range")
            label = (f" [dates {dr[0]}-{dr[1]}; may be outdated]"
                     if dr else " [may be outdated]")
            body = "\n\n".join(f"[scope={b['scope']}]\n{b['text']}"
                               for b in notam["blocks"])
            parts += [f"\n=== SECTION 4: NOTAMs{label} ===",
                      body[:BUDGET_NOTAM]]
    return "\n".join(parts)

# ----------------------------------------------------------------------------
# LLM extraction (batched, resumable)
# ----------------------------------------------------------------------------

def batch_system(batch):
    return ("You extract facts from an official aviation dossier (Foreign "
            "Clearance Guide plus airfield/support/NOTAM sections) for one "
            "country.\nSTRICT RULES:\n"
            "- Answer ONLY from the dossier text. If a fact is not stated, "
            "the answer is exactly \"NA\". Never use outside knowledge.\n"
            "- Each field asks one question; answer each independently with "
            "verbatim quotes from the dossier. Note which section a quote "
            "came from if ambiguous.\n"
            "Respond ONLY with a JSON object with EXACTLY these keys:\n"
            + json.dumps({n: "..." for n, _ in batch}) + "\n"
            "Field questions:\n"
            + "\n".join(f"- {n}: {q}" for n, q in batch))


def fields_hash(batch) -> str:
    return hashlib.sha256(
        json.dumps([n for n, _ in batch]).encode()).hexdigest()[:12]


_client = None
_sem = None


def _get_client():
    global _client
    if _client is None:
        import httpx
        from openai import AsyncOpenAI
        http_client = None if TLS_VERIFY else httpx.AsyncClient(verify=False)
        _client = AsyncOpenAI(base_url=BASE_URL, api_key=API_KEY,
                              http_client=http_client)
    return _client


def _get_sem():
    global _sem
    if _sem is None:
        _sem = asyncio.Semaphore(CONCURRENCY)
    return _sem


async def call_llm(system, user):
    async with _get_sem():
        for attempt in range(1, RETRIES + 1):
            try:
                r = await _get_client().chat.completions.create(
                    model=MODEL, temperature=0, max_tokens=MAX_TOKENS,
                    messages=[{"role": "system", "content": system},
                              {"role": "user", "content": user}])
                return r.choices[0].message.content or ""
            except Exception as e:  # noqa: BLE001
                print(f"  [llm] attempt {attempt} failed: {e}")
                await asyncio.sleep(3 * attempt)
    return ""


async def extract_batch(country, dossier, batch, b_idx):
    resp = parse_lenient(await call_llm(batch_system(batch),
                                        f"DOSSIER FOR {country}:\n\n{dossier}"))
    row = {"country": country, "batch": b_idx, "fhash": fields_hash(batch),
           "_status": "ok" if isinstance(resp, dict) else "llm_failure"}
    resp = resp if isinstance(resp, dict) else {}
    for name, _ in batch:
        v = resp.get(name, "NA")
        row[name] = "NA" if v is None or not str(v).strip() else str(v)
    return row


async def run(items_path, countries_dir, sources_dir, out_csv):
    fields = build_fields(load_items(items_path))
    batches = batch_fields(fields)
    field_names = [n for n, _ in fields]
    print(f"[fields] {len(fields)} total in {len(batches)} batches "
          f"of <= {FIELDS_PER_CALL}")

    country_files = sorted(Path(countries_dir).glob("*.json"))
    if not country_files:
        sys.exit(f"No country JSON files in {countries_dir}")

    slices = {}
    if sources_dir:
        from country_sources import build_or_load_slices, summary_table
        slices = build_or_load_slices(countries_dir, sources_dir,
                                      "data/work/sources_index")
        print(summary_table(slices))

    dossier_dir = Path("data/work/dossiers")
    dossier_dir.mkdir(parents=True, exist_ok=True)
    dossiers = {}
    for p in country_files:
        c = country_key(p)
        dossiers[c] = build_dossier(c, p, slices.get(c))
        (dossier_dir / f"{c}.txt").write_text(dossiers[c], encoding="utf-8")
    print(f"[dossiers] written to {dossier_dir} for audit "
          f"(sizes: {', '.join(f'{c}:{len(d)}' for c, d in dossiers.items())})")

    # Resume: (country, batch, fhash) triples already done
    ckpt_path = out_csv.with_suffix(".checkpoint.jsonl")
    done = {}
    if ckpt_path.exists():
        with open(ckpt_path, encoding="utf-8") as f:
            for line in f:
                if line.strip():
                    r = json.loads(line)
                    done[(r["country"], r["batch"], r.get("fhash"))] = r

    todo = []
    for c in dossiers:
        for b_idx, batch in enumerate(batches):
            if (c, b_idx, fields_hash(batch)) not in done:
                todo.append((c, b_idx, batch))
    total_units = len(dossiers) * len(batches)
    print(f"[run] {total_units} country-batches, "
          f"{total_units - len(todo)} cached, {len(todo)} to extract")

    with open(ckpt_path, "a", encoding="utf-8") as ckpt:
        tasks = [asyncio.create_task(
            extract_batch(c, dossiers[c], batch, b_idx))
            for c, b_idx, batch in todo]
        n_done = total_units - len(todo)
        for coro in asyncio.as_completed(tasks):
            row = await coro
            done[(row["country"], row["batch"], row["fhash"])] = row
            ckpt.write(json.dumps(row, ensure_ascii=False) + "\n")
            ckpt.flush()
            n_done += 1
            print(f"[{n_done}/{total_units}] {row['country']} "
                  f"batch {row['batch'] + 1}/{len(batches)} "
                  f"({row['_status']})")

    # Merge batches -> one row per country
    merged = {}
    fails = []
    for b_idx, batch in enumerate(batches):
        fh = fields_hash(batch)
        for c in dossiers:
            row = done.get((c, b_idx, fh))
            m = merged.setdefault(c, {"country": c, "_status": "ok"})
            if row is None or row["_status"] != "ok":
                m["_status"] = "partial"
                fails.append((c, b_idx))
                for name, _ in batch:
                    m.setdefault(name, "NA")
            else:
                for name, _ in batch:
                    m[name] = row.get(name, "NA")

    cols = ["country", "_status"] + field_names
    with open(out_csv, "w", newline="", encoding="utf-8") as f:
        w = csv.DictWriter(f, fieldnames=cols, extrasaction="ignore")
        w.writeheader()
        for c in sorted(merged):
            w.writerow(merged[c])
    print(f"\nWrote {len(merged)} rows x {len(field_names)} fields "
          f"-> {out_csv}")
    if fails:
        print(f"PARTIAL: {len(fails)} failed batches {fails} — delete their "
              f"lines from {ckpt_path.name} and rerun to retry.")
    print(f'Next: FCG_CSV={out_csv} python country_route_planner.py '
          f'"A to B to C"')


def main():
    ap = argparse.ArgumentParser()
    ap.add_argument("items", type=Path, help="items file (.json or .docx)")
    ap.add_argument("--countries", type=Path, default=Path("data/fcg_countries"))
    ap.add_argument("--sources", type=Path, default=None,
                    help="dir with meis/SVC_RMK/udl files (optional)")
    ap.add_argument("--out", type=Path, default=Path("fcg_extract.csv"))
    args = ap.parse_args()
    if not args.items.exists():
        sys.exit(f"Items file not found: {args.items}")
    asyncio.run(run(args.items, args.countries, args.sources, args.out))


if __name__ == "__main__":
    main()
